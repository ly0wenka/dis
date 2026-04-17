from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math2docx

# Створення документа
doc = Document()

# Налаштування стилю
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

# ============================================================================
# ЗАГОЛОВКИ
# ============================================================================

heading = doc.add_heading('Розділ 2', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_heading(
    'МЕТОДИ ДИСТАНЦІЙНОЇ ІДЕНТИФІКАЦІЇ НА ОСНОВІ КОМП\'ЮТЕРНОГО ЗОРУ ТА ГЛИБИННОГО НАВЧАННЯ',
    level=2
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================================
# ПІДРОЗДІЛ 2.1
# ============================================================================

doc.add_heading('2.1 Методи оптичного потоку в задачах ідентифікації', level=2)

doc.add_paragraph("""Алгоритм Горна — Шунка спирається на припущення про плавність потоку над усім зображенням. 
Потік формулюють як глобальний функціонал енергії:""")

# Формула (2.1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r'E = \iint \left( I_x u + I_y v + I_t \right)^2 + \alpha^2 \left( |\nabla u|^2 + |\nabla v|^2 \right) dx dy')
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.1)')

doc.add_paragraph("""де I_x, I_y, I_t — похідні інтенсивності за x, y та часом, 
V = [u(x,y), v(x,y)]^T — вектор оптичного потоку, α — стала регуляризації.""")

# Формули (2.2)-(2.3)
for eq, num in [
    (r'\frac{\partial L}{\partial u} - \frac{\partial}{\partial x}\frac{\partial L}{\partial u_x} - \frac{\partial}{\partial y}\frac{\partial L}{\partial u_y} = 0', '2.2'),
    (r'\frac{\partial L}{\partial v} - \frac{\partial}{\partial x}\frac{\partial L}{\partial v_x} - \frac{\partial}{\partial y}\frac{\partial L}{\partial v_y} = 0', '2.3')
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math2docx.add_math(p, eq)
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.add_run(f'({num})')

# Формули (2.4)-(2.5)
for eq, num in [
    (r'I_x(I_x u + I_y v + I_t) - \alpha^2 \Delta u = 0', '2.4'),
    (r'I_y(I_x u + I_y v + I_t) - \alpha^2 \Delta v = 0', '2.5')
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math2docx.add_math(p, eq)
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.add_run(f'({num})')

# Формула (2.6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r'\Delta u(x,y) = 4(\bar{u}(x,y) - u(x,y))')
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.6)')

# Формули (2.7)-(2.8)
for eq, num in [
    (r'(I_x^2 + 4\alpha^2)u + I_x I_y v = 4\alpha^2\bar{u} - I_x I_t', '2.7'),
    (r'I_x I_y u + (I_y^2 + 4\alpha^2)v = 4\alpha^2\bar{v} - I_y I_t', '2.8')
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math2docx.add_math(p, eq)
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.add_run(f'({num})')

doc.add_paragraph("Ітераційна схема методу Якобі:")

# Формули (2.9)-(2.10)
for eq, num in [
    (r'u^{k+1} = \bar{u}^k - \frac{I_x (I_x \bar{u}^k + I_y \bar{v}^k + I_t)}{4\alpha^2 + I_x^2 + I_y^2}', '2.9'),
    (r'v^{k+1} = \bar{v}^k - \frac{I_y (I_x \bar{u}^k + I_y \bar{v}^k + I_t)}{4\alpha^2 + I_x^2 + I_y^2}', '2.10')
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math2docx.add_math(p, eq)
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.add_run(f'({num})')

# ============================================================================
# МЕТОД ЛУКАСА-КАНАДЕ (2.11-2.13)
# ============================================================================

doc.add_heading('Метод Лукаса-Канаде', level=3)

for eq, num in [
    (r'I_x(q_i)V_x + I_y(q_i)V_y = -I_t(q_i)', '2.11'),
    (r'A v = b', '2.12'),
    (r'v = (A^T A)^{-1} A^T b', '2.13')
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math2docx.add_math(p, eq)
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.add_run(f'({num})')

# ============================================================================
# ПІДРОЗДІЛ 2.2 - GEONET
# ============================================================================

doc.add_heading('2.2 Нейромережеві моделі для аналізу руху', level=2)

# Формула (2.14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r'F_{t \rightarrow s}^{rig}(p_t) = K T_{t \rightarrow s} D_t(p_t) K^{-1} p_t - p_t')
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.14)')

# Фотометрична втрата (2.15)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r'\mathcal{L}_{rw} = \alpha \frac{1 - SSIM(I_t, \tilde{I}_s^{rig})}{2} + (1 - \alpha) \| I_t - \tilde{I}_s^{rig} \|_1')
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.15)')

# Гладкість глибини (2.16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r'\mathcal{L}_{ds} = \sum |\nabla D(p_t)| \cdot (e^{-|\nabla I(p_t)|})^T')
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.16)')

# Геометрична узгодженість (2.17)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r'\mathcal{L}_{gc} = \sum [\delta(p_t)] \cdot \| \Delta f_{t \rightarrow s}^{full}(p_t) \|_1')
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.17)')

# Повна втрата (2.18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r'\mathcal{L} = \sum_l \sum_{\langle t,s \rangle} \{ \mathcal{L}_{rw} + \lambda_{ds} \mathcal{L}_{ds} + \mathcal{L}_{fw} + \lambda_{fs} \mathcal{L}_{fs} + \lambda_{gc} \mathcal{L}_{gc} \}')
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.18)')

# ============================================================================
# ПІДРОЗДІЛ 2.4 - ОБ'ЄДНАННЯ МЕТОДІВ
# ============================================================================

doc.add_heading('2.4 Об\'єднання методів', level=2)

for eq, num in [
    (r'I(x,y,t): \Omega \subset \mathbb{R}^2 \times [0,T] \rightarrow \mathbb{R}', '2.19'),
    (r'X(t) = \{ x_i(t) \}_{i=1}^{N}', '2.20'),
    (r'S(t) = [x(t), y(t), v_x(t), v_y(t), \theta(t)]^T', '2.21'),
    (r'\dot{s}(t) = f(s(t), w(t))', '2.22')
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math2docx.add_math(p, eq)
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.add_run(f'({num})')

# ============================================================================
# ВИСНОВКИ
# ============================================================================

doc.add_heading('2.5 Висновки до розділу', level=2)

conclusions = [
    "1. Проаналізовано методи оптичного потоку: класичні (Горна-Шунка, Лукаса-Канаде) та нейромережеві (FlowNet, FlowNet 2.0).",
    "2. Досліджено архітектуру GeoNet, яка поєднує оцінку глибини та оптичного потоку в єдиній неконтрольованій структурі.",
    "3. Розглянуто ансамблеві методи (бегінг та бустінг) для підвищення стійкості ідентифікації.",
    "4. Запропоновано структуру об'єднання методів DETR, оптичного потоку та GeoNet."
]

for concl in conclusions:
    doc.add_paragraph(concl, style='List Bullet')

# ============================================================================
# ЗБЕРЕЖЕННЯ
# ============================================================================

doc.save('Rozdil_2_math2docx.docx')

print("=" * 70)
print("✅ Документ 'Rozdil_2_math2docx.docx' успішно створено!")
print("=" * 70)
print("\n📊 Додано формули (2.1) - (2.22) з використанням math2docx")