from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math2docx

# ============================================================================
# СТВОРЕННЯ ДОКУМЕНТА
# ============================================================================

doc = Document()

# Налаштування стилю
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ============================================================================
# ЗАГОЛОВОК
# ============================================================================

title = doc.add_heading('An Iterative Image Registration Technique with an Application to Stereo Vision', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Автори
authors = doc.add_paragraph('Bruce D. Lucas and Takeo Kanade')
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.runs[0].italic = True

# Афіліація
aff = doc.add_paragraph('Computer Science Department, Carnegie-Mellon University')
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Рік
year = doc.add_paragraph('1981')
year.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================================
# АНОТАЦІЯ (ABSTRACT)
# ============================================================================

doc.add_heading('Abstract', level=2)

abstract_text = """Image registration finds a variety of applications in computer vision. Unfortunately, traditional image registration techniques tend to be costly. We present a new image registration technique that makes use of the spatial intensity gradient of the images to find a good match using a type of Newton-Raphson iteration. Our technique is faster because it examines far fewer potential matches between the images than existing techniques. Furthermore, this registration technique can be generalized to handle rotation, scaling and shearing. We show how our technique can be adapted for use in a stereo vision system."""
doc.add_paragraph(abstract_text)

doc.add_page_break()

# ============================================================================
# 1. INTRODUCTION
# ============================================================================

doc.add_heading('1. Introduction', level=2)

intro_text = """Image registration finds a variety of applications in computer vision, such as image matching for stereo vision, pattern recognition, and motion analysis. Unfortunately, existing techniques for image registration tend to be costly. Moreover, they generally fail to deal with rotation or other distortions of the images.

In this paper we present a new image registration technique that uses spatial intensity gradient information to direct the search for the position that yields the best match. By taking more information about the images into account, this technique is able to find the best match between two images with far fewer comparisons of images than techniques which examine the possible positions of registration in some fixed order. Our technique takes advantage of the fact that in many applications the two images are already in approximate registration. This technique can be generalized to deal with arbitrary linear distortions of the image, including rotation."""
doc.add_paragraph(intro_text)

# ============================================================================
# 2. THE REGISTRATION PROBLEM
# ============================================================================

doc.add_heading('2. The Registration Problem', level=2)

reg_text = """The translational image registration problem can be characterized as follows: We are given functions F(x) and G(x) which give the respective pixel values at each location x in two images, where x is a vector. We wish to find the disparity vector h which minimizes some measure of the difference between F(x+h) and G(x), for x in some region of interest R.

Typical measures of the difference between F(x+h) and G(x) are:"""
doc.add_paragraph(reg_text)

doc.add_paragraph('L₁ norm = Σ|F(x+h) - G(x)|', style='List Bullet')
doc.add_paragraph('L₂ norm = √(Σ[F(x+h) - G(x)]²)', style='List Bullet')
doc.add_paragraph('Negative of normalized correlation = -ΣF(x+h)G(x) / √(ΣF(x+h)² ΣG(x)²)', style='List Bullet')

# ============================================================================
# 4. THE REGISTRATION ALGORITHM
# ============================================================================

doc.add_heading('4. The Registration Algorithm', level=2)

# ============================================================================
# 4.1. One Dimensional Case
# ============================================================================

doc.add_heading('4.1. One Dimensional Case', level=3)

doc.add_paragraph('For small h, the linear approximation gives:')

# Формула (1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"F'(x) \approx \frac{F(x+h)-F(x)}{h} = \frac{G(x)-F(x)}{h}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(1)')

doc.add_paragraph('Therefore:')

# Формула (2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h \approx \frac{G(x)-F(x)}{F'(x)}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2)')

doc.add_paragraph('Averaging multiple estimates:')

# Формула (3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \sum_x \frac{G(x)-F(x)}{F'(x)} / \sum_x 1")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3)')

doc.add_paragraph('Weighted average with w(x) = 1/|G\'(x)-F\'(x)|:')

# Формула (6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \sum_x \frac{w(x)[G(x)-F(x)]}{F'(x)} / \sum_x w(x)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(6)')

doc.add_paragraph('Newton-Raphson iteration:')

# Формула (7)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h_0 = 0,\quad h_{k+1} = h_k + \sum_x \frac{w(x)[G(x)-F(x+h_k)]}{F'(x+h_k)} / \sum_x w(x)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(7)')

# ============================================================================
# 4.2. Alternative Derivation (Minimizing L₂ Norm)
# ============================================================================

doc.add_heading('4.2. Alternative Derivation (Minimizing L₂ Norm)', level=3)

doc.add_paragraph('Linear approximation:')

# Формула (8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"F(x+h) \approx F(x) + hF'(x)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(8)')

doc.add_paragraph('Error function:')

# Формула (9) - перша частина
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E = \sum_x [F(x) + hF'(x) - G(x)]^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(9)')

doc.add_paragraph('Setting derivative to zero yields:')

# Формула (10)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \frac{\sum_x F'(x)[G(x)-F(x)]}{\sum_x F'(x)^2}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(10)')

# ============================================================================
# 4.4. Generalization to Multiple Dimensions
# ============================================================================

doc.add_heading('4.4. Generalization to Multiple Dimensions', level=3)

doc.add_paragraph('For n-dimensional case, we minimize:')

# Формула багатовимірна
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E = \sum_{x \in R} [F(x+h) - G(x)]^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(11)')

doc.add_paragraph('Linear approximation in n-dimensions:')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"F(x+h) \approx F(x) + h \frac{\partial F}{\partial x}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(12)')

doc.add_paragraph('The solution for h:')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \left[ \sum_x \left( \frac{\partial F}{\partial x} \right)^T (G(x)-F(x)) \right] \left[ \sum_x \left( \frac{\partial F}{\partial x} \right)^T \left( \frac{\partial F}{\partial x} \right) \right]^{-1}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(13)')

# ============================================================================
# 5. APPLICATION TO STEREO VISION
# ============================================================================

doc.add_heading('5. Application to Stereo Vision', level=2)

doc.add_paragraph("""
The generalized registration algorithm can be applied to extracting depth information from stereo images.
For an object at distance z from camera 1, the position in camera 2's film plane can be calculated.
Using the linear approximation for depth adjustment Δz:
""")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\Delta z = \sum_x \frac{\partial F}{\partial z} [G-F] / \sum_x \left( \frac{\partial F}{\partial z} \right)^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(14)')

doc.add_paragraph("""
where ∂F/∂z is computed via the chain rule from the spatial intensity gradient and the geometry of the stereo setup.
Similarly, for unknown camera parameters c, the update Δc is:
""")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\Delta c = \left[ \sum_x \left( \frac{\partial q}{\partial c} \frac{\partial F}{\partial q} \right)^T (G-F) \right] \left[ \sum_x \left( \frac{\partial q}{\partial c} \frac{\partial F}{\partial q} \right)^T \left( \frac{\partial q}{\partial c} \frac{\partial F}{\partial q} \right) \right]^{-1}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(15)')

# ============================================================================
# ВИСНОВКИ
# ============================================================================

doc.add_heading('Conclusions', level=2)

conclusions = """1. The proposed iterative image registration technique uses spatial intensity gradient information to guide the search, requiring far fewer comparisons than exhaustive search methods.

2. The method provides a Newton-Raphson style iteration that converges quadratically when the initial estimate is sufficiently close.

3. The technique generalizes naturally to multiple dimensions and to arbitrary linear transformations (rotation, scaling, shearing).

4. In stereo vision applications, the algorithm can simultaneously solve for object depths and camera parameters.

5. The coarse-fine strategy using bandpass-filtered images extends the range of convergence while maintaining accuracy."""
doc.add_paragraph(conclusions)

# ============================================================================
# ЗБЕРЕЖЕННЯ
# ============================================================================

output_file = 'Lucas_Kanade_1981.docx'
doc.save(output_file)

print("=" * 70)
print(f"✅ Документ '{output_file}' успішно створено!")
print("=" * 70)
print("\n📄 Документ містить:")
print("   • Повну назву статті та інформацію про авторів")
print("   • Анотацію (Abstract)")
print("   • Вступ (Introduction)")
print("   • Постановку задачі реєстрації зображень")
print("   • Одновимірний випадок методу Лукаса-Канаде (формули 1-7)")
print("   • Альтернативний висновок через мінімізацію L₂ норми (формули 8-10)")
print("   • Узагальнення на багатовимірний випадок (формули 11-13)")
print("   • Застосування до стереозору (формули 14-15)")
print("   • Висновки")