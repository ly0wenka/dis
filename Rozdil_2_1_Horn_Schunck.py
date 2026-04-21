from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math2docx

# ============================================================================
# СТВОРЕННЯ ДОКУМЕНТА
# ============================================================================

doc = Document()

# Налаштування стилю
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

# ============================================================================
# ЗАГОЛОВОК ПІДРОЗДІЛУ
# ============================================================================

doc.add_heading('2.1 Методи оптичного потоку в задачах ідентифікації', level=2)

# ============================================================================
# ОДИН СУЦІЛЬНИЙ АБЗАЦ З ТЕКСТОМ PDF
# ============================================================================

text = """
Метод Горна‑Шунка є глобальним методом обчислення оптичного потоку, представленим у 1981 році, який базується на двох основних припущеннях, а саме сталості яскравості, тобто колір або яскравість точки об'єкта не змінюється між кадрами, та плавності потоку, тобто оптичний потік змінюється плавно від пікселя до пікселя, і на відміну від локального методу Лукаса‑Канаде, який передбачає сталість потоку в локальному вікні, метод Горна‑Шунка є глобальним і дає щільне поле оптичного потоку. 
Сталість яскравості дозволяє порівнювати пікселі між кадрами за рівнянням I(x, y, t) = I(x + uδt, y + vδt, t + δt), а малий рух означає, що пікселі переміщуються лише на невелику відстань між кадрами, що дозволяє лінеаризувати обмеження сталості яскравості. 
Розклад в ряд Тейлора для малого δt дає I(x, y, t) + (∂I/∂x)δx + (∂I/∂y)δy + (∂I/∂t)δt = I(x, y, t), що після спрощення дає рівняння оптичного потоку I_x u + I_y v + I_t = 0, де I_x та I_y є просторовими градієнтами зображення, а I_t є часовою похідною. 
Метод Горна‑Шунка поєднує два обмеження, а саме сталість яскравості та плавність потоку, де складова сталості яскравості має вигляд E_d(i,j) = [I_x u_ij + I_y v_ij + I_t]², а складова плавності штрафує різкі зміни оптичного потоку між сусідніми пікселями і визначається як E_s(i,j) = 1/4[(u_ij - u_i+1,j)² + (u_ij - u_i,j+1)² + (v_ij - v_i+1,j)² + (v_ij - v_i,j+1)²]. 
Повний функціонал, який мінімізується, має вигляд min_{u,v} Σ_{i,j} {E_s(i,j) + λ E_d(i,j)}, де λ є параметром, що регулює вагу між гладкістю та відповідністю яскравості. 
Локальне середнє обчислюється як ū_ij = 1/4(u_i+1,j + u_i-1,j + u_i,j+1 + u_i,j-1) і аналогічно для v̄_ij. 
Часткові похідні функціоналу за u_kl та v_kl мають вигляд ∂E/∂u_kl = 2(u_kl - ū_kl) + 2λ(I_x u_kl + I_y v_kl + I_t)I_x та ∂E/∂v_kl = 2(v_kl - v̄_kl) + 2λ(I_x u_kl + I_y v_kl + I_t)I_y. 
Прирівнювання похідних до нуля дає систему лінійних рівнянь (1 + λ I_x²)u_kl + λ I_x I_y v_kl = ū_kl - λ I_x I_t та λ I_x I_y u_kl + (1 + λ I_y²)v_kl = v̄_kl - λ I_y I_t, яка є системою вигляду Ax = b і розв'язується для кожного пікселя. 
Після перетворень отримуємо {1 + λ(I_x² + I_y²)} u_kl = (1 + λ I_x²)ū_kl - λ I_x I_y v̄_kl - λ I_x I_t та {1 + λ(I_x² + I_y²)} v_kl = (1 + λ I_y²)v̄_kl - λ I_x I_y ū_kl - λ I_y I_t. 
Звідси виводиться ітераційна схема оновлення у вигляді \hat{u}_kl = ū_kl - (I_x ū_kl + I_y v̄_kl + I_t) I_x / (λ⁻¹ + I_x² + I_y²) та \hat{v}_kl = v̄_kl - (I_x ū_kl + I_y v̄_kl + I_t) I_y / (λ⁻¹ + I_x² + I_y²). 
Параметр λ контролює баланс між двома складовими функціоналу, де при малому λ домінує складова плавності, а при великому λ домінує складова сталості яскравості, і типове значення λ знаходиться в діапазоні від 0,1 до 1,0. 
Алгоритм Горна‑Шунка складається з наступних кроків: спочатку обчислюються просторові градієнти зображення I_x та I_y, потім обчислюється часовий градієнт I_t як різниця між кадрами, далі ініціалізується поле оптичного потоку u = 0 та v = 0, і після цього повторюється до збіжності обчислення локального середнього ū та v̄ та оновлення u та v за наведеними ітераційними формулами. 
На практиці достатньо від 50 до 100 ітерацій для досягнення збіжності. 
Порівняно з методом Лукаса‑Канаде, метод Горна‑Шунка є глобальним методом, тоді як Лукаса‑Канаде є локальним методом, він використовує припущення про плавність потоку, тоді як Лукаса‑Канаде використовує припущення про сталість потоку у вікні, він дає щільне поле потоку, а Лукаса‑Канаде дає розріджене поле потоку, він є ітераційним, тоді як Лукаса‑Канаде є аналітичним, і він має вищу стійкість до шуму, ніж Лукаса‑Канаде. 
Метод Горна‑Шунка є фундаментальним для розуміння задачі оптичного потоку та основою для багатьох сучасних алгоритмів, однак його недоліками є чутливість до порушення припущення про сталість яскравості, наприклад при тінях або змінах освітлення, та розмиття меж об'єктів через квадратичну регуляризацію.
"""

doc.add_paragraph(text)

# ============================================================================
# ДОДАВАННЯ ФОРМУЛ
# ============================================================================

doc.add_heading('Основні формули методу Горна‑Шунка', level=3)

# Формула (3.17) - рівняння сталості яскравості
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"I(x + u\delta t, y + v\delta t, t + \delta t) = I(x, y, t)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.17)')

# Формула (3.18) - розклад в ряд Тейлора
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"I(x, y, t) + \frac{\partial I}{\partial x}\delta x + \frac{\partial I}{\partial y}\delta y + \frac{\partial I}{\partial t}\delta t = I(x, y, t)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.18)')

# Формула (3.19) - рівняння оптичного потоку
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"I_x u + I_y v + I_t = 0")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.19)')

# Формула (3.20) - складова сталості яскравості
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E_d(i,j) = \left[ I_x u_{ij} + I_y v_{ij} + I_t \right]^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.20)')

# Формула (3.21) - повний функціонал
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\min_{u,v} \sum_{i,j} \left\{ E_s(i,j) + \lambda E_d(i,j) \right\}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.21)')

# Формула (3.22) - локальне середнє
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\bar{u}_{ij} = \frac{1}{4} \left( u_{i+1,j} + u_{i-1,j} + u_{i,j+1} + u_{i,j-1} \right)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.22)')

# Формула (3.23) - часткова похідна за u
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\frac{\partial E}{\partial u_{kl}} = 2(u_{kl} - \bar{u}_{kl}) + 2\lambda (I_x u_{kl} + I_y v_{kl} + I_t) I_x")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.23)')

# Формула (3.24) - часткова похідна за v
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\frac{\partial E}{\partial v_{kl}} = 2(v_{kl} - \bar{v}_{kl}) + 2\lambda (I_x u_{kl} + I_y v_{kl} + I_t) I_y")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.24)')

# Формула (3.25) - лінійна система
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"(1 + \lambda I_x^2) u_{kl} + \lambda I_x I_y v_{kl} = \bar{u}_{kl} - \lambda I_x I_t")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.25)')

# Формула (3.26) - розв'язок для u
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\{1 + \lambda(I_x^2 + I_y^2)\} u_{kl} = (1 + \lambda I_x^2) \bar{u}_{kl} - \lambda I_x I_y \bar{v}_{kl} - \lambda I_x I_t")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.26)')

# Формула (3.27) - розв'язок для v
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\{1 + \lambda(I_x^2 + I_y^2)\} v_{kl} = (1 + \lambda I_y^2) \bar{v}_{kl} - \lambda I_x I_y \bar{u}_{kl} - \lambda I_y I_t")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.27)')

# Формула (3.28) - ітераційна схема для u
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\hat{u}_{kl} = \bar{u}_{kl} - \frac{I_x \bar{u}_{kl} + I_y \bar{v}_{kl} + I_t}{\lambda^{-1} + I_x^2 + I_y^2} I_x")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.28)')

# Формула (3.29) - ітераційна схема для v
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\hat{v}_{kl} = \bar{v}_{kl} - \frac{I_x \bar{u}_{kl} + I_y \bar{v}_{kl} + I_t}{\lambda^{-1} + I_x^2 + I_y^2} I_y")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.29)')

# Формула (3.30) - складова плавності
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E_s(i,j) = \frac{1}{4} \left[ (u_{ij} - u_{i+1,j})^2 + (u_{ij} - u_{i,j+1})^2 + (v_{ij} - v_{i+1,j})^2 + (v_{ij} - v_{i,j+1})^2 \right]")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3.30)')

# ============================================================================
# ЗБЕРЕЖЕННЯ
# ============================================================================

output_file = 'Rozdil_2_1_Horn_Schunck.docx'
doc.save(output_file)

print("=" * 70)
print(f"✅ Документ '{output_file}' успішно створено!")
print("=" * 70)
print("\n📄 Документ містить:")
print("   • Підрозділ 2.1 – Методи оптичного потоку в задачах ідентифікації")
print("   • Один суцільний абзац тексту з описом методу Горна‑Шунка")
print("   • 13 основних формул (3.17) – (3.30)")