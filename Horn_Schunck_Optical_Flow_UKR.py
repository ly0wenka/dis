from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math2docx

# ============================================================================
# СТВОРЕННЯ ДОКУМЕНТА
# ============================================================================

doc = Document()

# Налаштування стилю
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

# ============================================================================
# ЗАГОЛОВОК
# ============================================================================

title = doc.add_heading('Метод Горна-Шунка для обчислення оптичного потоку', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

authors = doc.add_paragraph('Кріс Кітані (Carnegie Mellon University)')
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.runs[0].italic = True

doc.add_paragraph()

# ============================================================================
# ВСТУП
# ============================================================================

doc.add_heading('Вступ', level=2)

intro_text = """Метод Горна-Шунка (Horn-Schunck) є глобальним методом обчислення оптичного потоку, представленим у 1981 році. Він базується на двох основних припущеннях:

1. Сталість яскравості (Brightness constancy) — колір/яскравість точки об'єкта не змінюється між кадрами
2. Плавність потоку (Smooth flow) — оптичний потік змінюється плавно від пікселя до пікселя

На відміну від локального методу Лукаса-Канаде, який передбачає постійність потоку в локальному вікні, метод Горна-Шунка є глобальним і дає щільне поле оптичного потоку."""
doc.add_paragraph(intro_text)

# ============================================================================
# ОСНОВНІ ПРИПУЩЕННЯ
# ============================================================================

doc.add_heading('Основні припущення', level=2)

doc.add_paragraph("""
1. Сталість яскравості (Brightness constancy):
   - Дозволяє порівнювати пікселі між кадрами
   - I(x, y, t) = I(x + uδt, y + vδt, t + δt)

2. Малий рух (Small motion):
   - Пікселі переміщуються лише на невелику відстань між кадрами
   - Дозволяє лінеаризувати обмеження сталості яскравості
""")

# ============================================================================
# ОБМЕЖЕННЯ СТАЛОСТІ ЯСКРАВОСТІ
# ============================================================================

doc.add_heading('Обмеження сталості яскравості', level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"I(x + u\delta t, y + v\delta t, t + \delta t) = I(x, y, t)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(1)')

doc.add_paragraph("Розклад в ряд Тейлора для малого δt:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"I(x, y, t) + \frac{\partial I}{\partial x}\delta x + \frac{\partial I}{\partial y}\delta y + \frac{\partial I}{\partial t}\delta t = I(x, y, t)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2)')

doc.add_paragraph("Спрощення дає рівняння оптичного потоку:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"I_x u + I_y v + I_t = 0")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3)')

doc.add_paragraph("де I_x, I_y — просторові градієнти зображення, I_t — часова похідна.")

# ============================================================================
# ФУНКЦІОНАЛ ГОРНА-ШУНКА
# ============================================================================

doc.add_heading('Функціонал Горна-Шунка', level=2)

doc.add_paragraph("Метод поєднує два обмеження: сталість яскравості та плавність потоку.")

# Складова сталості яскравості
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E_d(i,j) = \left[ I_x u_{ij} + I_y v_{ij} + I_t \right]^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(4)')

# Складова плавності
doc.add_paragraph("Складова плавності (smoothness) — штрафує різкі зміни оптичного потоку між сусідніми пікселями:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E_s(i,j) = \frac{1}{4} \left[ (u_{ij} - u_{i+1,j})^2 + (u_{ij} - u_{i,j+1})^2 + (v_{ij} - v_{i+1,j})^2 + (v_{ij} - v_{i,j+1})^2 \right]")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(5)')

# Повний функціонал
doc.add_paragraph("Повний функціонал, який мінімізується:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\min_{u,v} \sum_{i,j} \left\{ E_s(i,j) + \lambda E_d(i,j) \right\}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(6)')

doc.add_paragraph("де λ — параметр, що регулює вагу між гладкістю та відповідністю яскравості.")

# ============================================================================
# ЛОКАЛЬНЕ СЕРЕДНЄ
# ============================================================================

doc.add_heading('Локальне середнє', level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\bar{u}_{ij} = \frac{1}{4} \left( u_{i+1,j} + u_{i-1,j} + u_{i,j+1} + u_{i,j-1} \right)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(7)')

doc.add_paragraph("Аналогічно для v̄_{ij}.")

# ============================================================================
# ЧАСТКОВІ ПОХІДНІ
# ============================================================================

doc.add_heading('Часткові похідні функціоналу', level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\frac{\partial E}{\partial u_{kl}} = 2(u_{kl} - \bar{u}_{kl}) + 2\lambda (I_x u_{kl} + I_y v_{kl} + I_t) I_x")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(8)')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\frac{\partial E}{\partial v_{kl}} = 2(v_{kl} - \bar{v}_{kl}) + 2\lambda (I_x u_{kl} + I_y v_{kl} + I_t) I_y")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(9)')

# ============================================================================
# ЛІНІЙНА СИСТЕМА
# ============================================================================

doc.add_heading('Лінійна система для екстремумів', level=2)

doc.add_paragraph("Прирівнювання похідних до нуля дає систему лінійних рівнянь:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"(1 + \lambda I_x^2) u_{kl} + \lambda I_x I_y v_{kl} = \bar{u}_{kl} - \lambda I_x I_t")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(10)')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\lambda I_x I_y u_{kl} + (1 + \lambda I_y^2) v_{kl} = \bar{v}_{kl} - \lambda I_y I_t")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(11)')

doc.add_paragraph("Це система вигляду Ax = b, яка розв'язується для кожного пікселя.")

# ============================================================================
# РОЗВ'ЯЗОК СИСТЕМИ
# ============================================================================

doc.add_heading('Розв\'язок системи', level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\{1 + \lambda(I_x^2 + I_y^2)\} u_{kl} = (1 + \lambda I_x^2) \bar{u}_{kl} - \lambda I_x I_y \bar{v}_{kl} - \lambda I_x I_t")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(12)')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\{1 + \lambda(I_x^2 + I_y^2)\} v_{kl} = (1 + \lambda I_y^2) \bar{v}_{kl} - \lambda I_x I_y \bar{u}_{kl} - \lambda I_y I_t")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(13)')

# ============================================================================
# ІТЕРАЦІЙНІ ФОРМУЛИ
# ============================================================================

doc.add_heading('Ітераційні формули оновлення', level=2)

doc.add_paragraph("Після перетворень отримуємо ітераційну схему:")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\hat{u}_{kl} = \bar{u}_{kl} - \frac{I_x \bar{u}_{kl} + I_y \bar{v}_{kl} + I_t}{\lambda^{-1} + I_x^2 + I_y^2} I_x")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(14)')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\hat{v}_{kl} = \bar{v}_{kl} - \frac{I_x \bar{u}_{kl} + I_y \bar{v}_{kl} + I_t}{\lambda^{-1} + I_x^2 + I_y^2} I_y")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(15)')

# ============================================================================
# АЛГОРИТМ
# ============================================================================

doc.add_heading('Алгоритм Горна-Шунка', level=2)

algorithm_text = """1. Обчислити просторові градієнти зображення I_x, I_y
2. Обчислити часовий градієнт I_t (різницю між кадрами)
3. Ініціалізувати поле оптичного потоку u = 0, v = 0
4. Повторювати до збіжності:
   Для кожного пікселя (i,j) обчислити локальне середнє ū, v̄
   Оновити u, v за формулами (14)-(15)

Алгоритм є ітераційним методом Якобі для розв'язання системи лінійних рівнянь.
На практиці достатньо 50-100 ітерацій для досягнення збіжності.
"""
doc.add_paragraph(algorithm_text)

# ============================================================================
# ІНТЕРПРЕТАЦІЯ ПАРАМЕТРА λ
# ============================================================================

doc.add_heading('Інтерпретація параметра λ', level=2)

doc.add_paragraph("""
Параметр λ контролює баланс між двома складовими функціоналу:

- Коли λ мале (λ⁻¹ велике): домінує складова плавності (smoothness)
- Коли λ велике: домінує складова сталості яскравості (brightness constancy)

Типове значення λ = 0.1 ... 1.0.
""")

# ============================================================================
# ПОРІВНЯННЯ З МЕТОДОМ ЛУКАСА-КАНАДЕ
# ============================================================================

doc.add_heading('Порівняння з методом Лукаса-Канаде', level=2)

comparison = """
| Характеристика | Горна-Шунка | Лукаса-Канаде |
|----------------|-------------|---------------|
| Тип методу | Глобальний | Локальний |
| Припущення | Плавність потоку | Постійність потоку у вікні |
| Результат | Щільне поле потоку | Розріджене поле потоку |
| Обчислення | Ітераційне | Аналітичне |
| Стійкість до шуму | Вища | Нижча |
"""
doc.add_paragraph(comparison)

# ============================================================================
# ВИСНОВКИ
# ============================================================================

doc.add_heading('Висновки', level=2)

conclusions = """1. Метод Горна-Шунка є глобальним підходом до обчислення оптичного потоку, що дає щільне поле векторів руху.

2. Основні припущення: сталість яскравості та плавність потоку.

3. Мінімізація функціоналу (6) призводить до ітераційної схеми (14)-(15).

4. Алгоритм є простим у реалізації (близько 8 рядків коду для основного циклу).

5. Недоліки: чутливість до порушення припущення про сталість яскравості (тіні, зміни освітлення) та розмиття меж об'єктів.

6. Метод є фундаментальним для розуміння задачі оптичного потоку та основою для багатьох сучасних алгоритмів."""
doc.add_paragraph(conclusions)

# ============================================================================
# ЗБЕРЕЖЕННЯ
# ============================================================================

output_file = 'Horn_Schunck_Optical_Flow_UKR.docx'
doc.save(output_file)

print("=" * 70)
print(f"✅ Документ '{output_file}' успішно створено українською мовою!")
print("=" * 70)
print("\n📄 Документ містить:")
print("   • Вступ та основні припущення методу Горна-Шунка")
print("   • Рівняння сталості яскравості (формули 1-3)")
print("   • Функціонал Горна-Шунка (формули 4-6)")
print("   • Локальне середнє (формула 7)")
print("   • Часткові похідні (формули 8-9)")
print("   • Лінійну систему (формули 10-11)")
print("   • Розв'язок системи (формули 12-13)")
print("   • Ітераційні формули оновлення (формули 14-15)")
print("   • Алгоритм Горна-Шунка")
print("   • Інтерпретацію параметра λ")
print("   • Порівняння з методом Лукаса-Канаде")
print("   • Висновки")