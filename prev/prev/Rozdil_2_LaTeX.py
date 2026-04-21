from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_eq.equation import Equation

doc = Document()

# Заголовок
heading = doc.add_heading('Розділ 2', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Формула (2.1)
eq = Equation(doc, r'E = \iint \left( I_x u + I_y v + I_t \right)^2 + \alpha^2 \left( |\nabla u|^2 + |\nabla v|^2 \right) dx dy')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run()._element.append(eq._element)

# Номер формули
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.1)')

# Формула (2.9) - ітераційна схема
eq2 = Equation(doc, r'u^{k+1} = \bar{u}^k - \frac{I_x (I_x \bar{u}^k + I_y \bar{v}^k + I_t)}{4\alpha^2 + I_x^2 + I_y^2}')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run()._element.append(eq2._element)
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2.9)')

doc.save('Rozdil_2_LaTeX.docx')
print("Документ з LaTeX-формулами створено!")