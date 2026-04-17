from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Створення нового документа
doc = Document()

# ============================================================================
# НАЛАШТУВАННЯ СТИЛІВ
# ============================================================================

# Стиль для формул
try:
    style = doc.styles.add_style('Formula', WD_STYLE_TYPE.PARAGRAPH)
except:
    style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
style.paragraph_format.space_before = Pt(6)
style.paragraph_format.space_after = Pt(6)

# Стиль для підписів рисунків
try:
    fig_style = doc.styles.add_style('FigureCaption', WD_STYLE_TYPE.PARAGRAPH)
except:
    fig_style = doc.styles['Normal']
fig_style.font.name = 'Times New Roman'
fig_style.font.size = Pt(12)
fig_style.font.italic = True
fig_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_style.paragraph_format.space_before = Pt(6)

# ============================================================================
# ЗАГОЛОВОК РОЗДІЛУ 2
# ============================================================================

heading = doc.add_heading('Розділ 2', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Додавання повної назви розділу
title = doc.add_paragraph()
title_run = title.add_run('МЕТОДИ ДИСТАНЦІЙНОЇ ІДЕНТИФІКАЦІЇ НА ОСНОВІ КОМП\'ЮТЕРНОГО ЗОРУ ТА ГЛИБИННОГО НАВЧАННЯ')
title_run.bold = True
title_run.font.size = Pt(14)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================================
# ПІДРОЗДІЛ 2.1
# ============================================================================

doc.add_heading('2.1 Методи оптичного потоку в задачах ідентифікації', level=2)

# Текст підрозділу
text_2_1 = """Алгоритм Горна — Шунка спирається на припущення про плавність потоку над усім зображенням. Таким чином, він намагається мінімізувати спотворення потоку й віддає перевагу розв'язкам, які демонструють більшу плавність.

Потік формулюють як глобальний функціонал енергії, який потім намагаються мінімізувати. Цю функцію для потоків двовимірних зображень задають як"""
doc.add_paragraph(text_2_1)

# Додавання формули (2.1)
p = doc.add_paragraph(style='Formula')
p.add_run('E = ∬[ (I_x u + I_y v + I_t)² + α²(|∇u|² + |∇v|²) ] dx dy')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.1)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Продовження тексту
text_2_1b = """де I_x, I_y, I_t — це похідні значень інтенсивності зображення за вимірами x, y та часу відповідно, V = [u(x,y), v(x,y)]ᵀ — це вектор оптичного потоку, а параметр α — це стала регуляризації. Більші значення α призводять до плавнішого потоку. Цей функціонал можливо мінімізувати, розв'язуючи пов'язані багатовимірні рівняння Ейлера — Лагранжа."""
doc.add_paragraph(text_2_1b)

# Рівняння (2.2) та (2.3)
p = doc.add_paragraph(style='Formula')
p.add_run('∂L/∂u - ∂/∂x·∂L/∂u_x - ∂/∂y·∂L/∂u_y = 0')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.2)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph(style='Formula')
p.add_run('∂L/∂v - ∂/∂x·∂L/∂v_x - ∂/∂y·∂L/∂v_y = 0')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.3)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Рівняння (2.4) та (2.5)
p = doc.add_paragraph(style='Formula')
p.add_run('I_x(I_x u + I_y v + I_t) - α²Δu = 0')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.4)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph(style='Formula')
p.add_run('I_y(I_x u + I_y v + I_t) - α²Δv = 0')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.5)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Рівняння (2.6)
p = doc.add_paragraph(style='Formula')
p.add_run('Δu(x,y) = 4(ū(x,y) - u(x,y))')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.6)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Лінійна система (2.7)-(2.8)
text_2_1c = """Використовуючи це позначення, наведену вище систему рівнянь можна записати як лінійну систему:"""
doc.add_paragraph(text_2_1c)

p = doc.add_paragraph(style='Formula')
p.add_run('(I_x² + 4α²)u + I_x I_y v = 4α²ū - I_x I_t')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.7)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph(style='Formula')
p.add_run('I_x I_y u + (I_y² + 4α²)v = 4α²v̄ - I_y I_t')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.8)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Ітераційна схема (2.9)-(2.10)
text_2_1d = """Наступна ітераційна схема виводиться за допомогою правила Крамера:"""
doc.add_paragraph(text_2_1d)

p = doc.add_paragraph(style='Formula')
p.add_run('uᵏ⁺¹ = ūᵏ - I_x(I_xūᵏ + I_yv̄ᵏ + I_t) / (4α² + I_x² + I_y²)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.9)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph(style='Formula')
p.add_run('vᵏ⁺¹ = v̄ᵏ - I_y(I_xūᵏ + I_yv̄ᵏ + I_t) / (4α² + I_x² + I_y²)')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.10)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ============================================================================
# ПІДРОЗДІЛ 2.2
# ============================================================================

doc.add_heading('2.2 Нейромережеві моделі для аналізу руху', level=2)

# Рівняння GeoNet (2.1 в розділі 2.2)
p = doc.add_paragraph(style='Formula')
p.add_run('F_{t→s}^{rig}(p_t) = K T_{t→s} D_t(p_t) K⁻¹ p_t - p_t')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.11)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Фотометрична втрата (2.2)
p = doc.add_paragraph(style='Formula')
p.add_run('L_rw = α·(1 - SSIM(I_t, Ĩ_s^{rig}))/2 + (1-α)·‖I_t - Ĩ_s^{rig}‖₁')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.12)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Гладкість глибини (2.3)
p = doc.add_paragraph(style='Formula')
p.add_run('L_ds = Σ|∇D(p_t)| · (e^{-|∇I(p_t)|})ᵀ')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.13)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Геометрична узгодженість (2.4)-(2.5)
p = doc.add_paragraph(style='Formula')
p.add_run('L_gc = Σ[δ(p_t)] · ‖Δf_{t→s}^{full}(p_t)‖₁')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.14)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Повна втрата (2.6)
p = doc.add_paragraph(style='Formula')
p.add_run('L = Σ_l Σ_{⟨t,s⟩} { L_rw + λ_ds L_ds + L_fw + λ_fs L_fs + λ_gc L_gc }')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.15)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Текст про генеративні моделі
text_2_2 = """Генеративні моделі, такі як StableAnimator та HunyuanVideo, вносять вагомий внесок у моделювання динамічних процесів. StableAnimator використовує дифузійні моделі та розв'язання рівняння Гамільтона-Якобі-Беллмана для створення високоточних анімацій із збереженням ідентичності. Метод SynCamMaster забезпечує синхронізовану генерацію відео з різних точок огляду, що є базою для створення «цифрових двійників» складних систем."""
doc.add_paragraph(text_2_2)

# ============================================================================
# ПІДРОЗДІЛ 2.3
# ============================================================================

doc.add_heading('2.3 Ансамблеві методи', level=2)

text_2_3 = """Ансамблевий метод Бегінг (bagging) полягає у створенні кількох моделей на основі різних підвибірок вхідних даних, об'єднаних за допомогою середнього значення або голосування. Суть цього методу полягає у створенні кількох моделей, кожна з яких тренується на випадковій підвибірці вхідних даних (кадрів відео). Це дозволяє зменшити дисперсію моделі та знизити ймовірність перенавчання.

Ансамблевий метод Бустінг (boosting) передбачає послідовне навчання моделей, де кожна наступна фокусується на помилках попередньої. Це дозволяє суттєво зменшити зміщення та підвищити точність прогнозування, особливо на межах об'єктів або в умовах слабкої освітленості."""
doc.add_paragraph(text_2_3)

# ============================================================================
# ПІДРОЗДІЛ 2.4
# ============================================================================

doc.add_heading('2.4 Об\'єднання методів', level=2)

# Просторово-часова модель (2.1 в розділі 2.4)
p = doc.add_paragraph(style='Formula')
p.add_run('I(x,y,t): Ω ⊂ ℝ² × [0,T] → ℝ')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.16)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Модель динамічного об'єкта (2.2)
p = doc.add_paragraph(style='Formula')
p.add_run('X(t) = { x_i(t) }_{i=1}^{N}')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.17)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Вектор стану (2.3)
p = doc.add_paragraph(style='Formula')
p.add_run('S(t) = [x(t), y(t), v_x(t), v_y(t), θ(t)]ᵀ')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.18)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Динамічна модель (2.4)
p = doc.add_paragraph(style='Formula')
p.add_run('ṡ(t) = f(s(t), w(t))')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Formula')
p.add_run('(2.19)')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ============================================================================
# ПІДРОЗДІЛ 2.5 - ВИСНОВКИ
# ============================================================================

doc.add_heading('2.5 Висновки до розділу', level=2)

conclusions = """1. Проаналізовано методи оптичного потоку: класичні (Горна-Шунка, Лукаса-Канаде) та нейромережеві (FlowNet, FlowNet 2.0). Визначено, що нейромережеві методи забезпечують вищу точність та деталізацію оцінки руху.

2. Досліджено архітектуру GeoNet, яка поєднує оцінку глибини та оптичного потоку в єдиній неконтрольованій структурі. Запропоновано модифікацію з додаванням трансформера для моделювання просторово-часових залежностей.

3. Розглянуто ансамблеві методи (бегінг та бустінг) для підвищення стійкості ідентифікації параметрів динамічних об'єктів в умовах шумів та оклюзій.

4. Запропоновано структуру об'єднання методів DETR, оптичного потоку та GeoNet для комплексної дистанційної ідентифікації параметрів динамічних об'єктів у відеопотоці."""
doc.add_paragraph(conclusions)

# ============================================================================
# ДОДАВАННЯ ТАБЛИЦІ
# ============================================================================

doc.add_heading('Порівняльна таблиця методів оптичного потоку', level=3)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'

# Заголовки таблиці
headers = ['Метод', 'Тип', 'Швидкодія', 'Точність (EPE)']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

# Дані таблиці
data = [
    ['Horn-Schunck', 'Глобальний', 'Низька', '~0.5-1.0 пікс'],
    ['Lucas-Kanade', 'Локальний', 'Висока', '~0.3-0.8 пікс'],
    ['Farneback', 'Поліноміальний', 'Середня', '~0.2-0.5 пікс'],
    ['FlowNet 2.0', 'Нейромережевий', 'Висока (140 fps)', '~0.1-0.3 пікс']
]

for i, row_data in enumerate(data, start=1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

# ============================================================================
# ДОДАВАННЯ СПИСКУ ЛІТЕРАТУРИ ДО РОЗДІЛУ
# ============================================================================

doc.add_heading('Література до розділу 2', level=3)

references = [
    'Yin Z., Shi J. GeoNet: Unsupervised Learning of Dense Depth, Optical Flow and Camera Pose. arXiv:1803.02276, 2018.',
    'Vaswani A., et al. Attention Is All You Need. NeurIPS, 2017.',
    'Dosovitskiy A., et al. FlowNet: Learning Optical Flow with Convolutional Networks. ICCV, 2015.',
    'Ilg E., et al. FlowNet 2.0: Evolution of Optical Flow Estimation with Deep Networks. CVPR, 2017.',
    'Carion N., et al. End-to-End Object Detection with Transformers. ECCV, 2020.'
]

for ref in references:
    p = doc.add_paragraph(ref, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

# ============================================================================
# ЗБЕРЕЖЕННЯ ДОКУМЕНТА
# ============================================================================

doc.save('Rozdil_2_Povnyi.docx')

print("=" * 60)
print("Документ 'Rozdil_2_Povnyi.docx' успішно створено!")
print("=" * 60)
print("\nЗміст документа:")
print("  • Розділ 2: Методи дистанційної ідентифікації...")
print("  • Підрозділ 2.1: Методи оптичного потоку (формули 2.1-2.10)")
print("  • Підрозділ 2.2: Нейромережеві моделі (формули 2.11-2.15)")
print("  • Підрозділ 2.3: Ансамблеві методи")
print("  • Підрозділ 2.4: Об'єднання методів (формули 2.16-2.19)")
print("  • Підрозділ 2.5: Висновки")
print("  • Таблиця порівняння методів оптичного потоку")
print("  • Список літератури")