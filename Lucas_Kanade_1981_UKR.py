from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math2docx

# ============================================================================
# СТВОРЕННЯ ДОКУМЕНТА
# ============================================================================

doc = Document()

# Налаштування стилю
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

# ============================================================================
# ЗАГОЛОВОК
# ============================================================================

title = doc.add_heading('Ітеративний метод реєстрації зображень та його застосування в стереобаченні', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Автори
authors = doc.add_paragraph('Брюс Д. Лукас та Такео Канаде')
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.runs[0].italic = True

# Афіліація
aff = doc.add_paragraph('Кафедра комп\'ютерних наук, Університет Карнегі-Меллон')
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Рік
year = doc.add_paragraph('1981')
year.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================================
# АНОТАЦІЯ (ABSTRACT)
# ============================================================================

doc.add_heading('Анотація', level=2)

abstract_text = """Реєстрація зображень має широкий спектр застосувань у комп'ютерному зорі. На жаль, традиційні методи реєстрації зображень є обчислювально затратними. Ми представляємо новий метод реєстрації зображень, який використовує просторовий градієнт інтенсивності зображень для знаходження відповідності за допомогою ітерацій Ньютона-Рафсона. Наш метод є швидшим, оскільки він досліджує значно менше потенційних відповідностей між зображеннями, ніж існуючі методи. Крім того, цей метод реєстрації може бути узагальнений для роботи з обертанням, масштабуванням та зсувом. Ми показуємо, як наш метод може бути адаптований для використання в системах стереобачення."""
doc.add_paragraph(abstract_text)

doc.add_page_break()

# ============================================================================
# 1. ВСТУП
# ============================================================================

doc.add_heading('1. Вступ', level=2)

intro_text = """Реєстрація зображень знаходить різноманітні застосування в комп'ютерному зорі, такі як зіставлення зображень для стереобачення, розпізнавання образів та аналіз руху. На жаль, існуючі методи реєстрації зображень є обчислювально затратними. Більше того, вони зазвичай не справляються з обертанням або іншими спотвореннями зображень.

У цій статті ми представляємо новий метод реєстрації зображень, який використовує інформацію про просторовий градієнт інтенсивності для спрямування пошуку позиції, що дає найкращу відповідність. Беручи до уваги більше інформації про зображення, цей метод здатний знайти найкращу відповідність між двома зображеннями з значно меншою кількістю порівнянь зображень, ніж методи, які досліджують можливі позиції реєстрації у фіксованому порядку. Наш метод використовує той факт, що в багатьох застосуваннях два зображення вже є приблизно зареєстрованими. Цей метод може бути узагальнений для роботи з довільними лінійними спотвореннями зображення, включаючи обертання."""
doc.add_paragraph(intro_text)

# ============================================================================
# 2. ЗАДАЧА РЕЄСТРАЦІЇ
# ============================================================================

doc.add_heading('2. Задача реєстрації', level=2)

reg_text = """Задачу реєстрації зображень при поступальному русі можна охарактеризувати наступним чином: Нехай задано функції F(x) та G(x), які визначають значення пікселів у відповідних позиціях x у двох зображеннях, де x є вектором. Ми хочемо знайти вектор зміщення h, який мінімізує деяку міру різниці між F(x+h) та G(x) для x у деякій області інтересу R.

Типовими мірами різниці між F(x+h) та G(x) є:"""
doc.add_paragraph(reg_text)

doc.add_paragraph('L₁ норма = Σ|F(x+h) - G(x)|', style='List Bullet')
doc.add_paragraph('L₂ норма = √(Σ[F(x+h) - G(x)]²)', style='List Bullet')
doc.add_paragraph('Від\'ємна нормована кореляція = -ΣF(x+h)G(x) / √(ΣF(x+h)² ΣG(x)²)', style='List Bullet')

# ============================================================================
# 4. МЕТОД РЕЄСТРАЦІЇ
# ============================================================================

doc.add_heading('4. Метод реєстрації', level=2)

# ============================================================================
# 4.1. Одновимірний випадок
# ============================================================================

doc.add_heading('4.1. Одновимірний випадок', level=3)

doc.add_paragraph('Для малих h лінійна апроксимація дає:')

# Формула (1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"F'(x) \approx \frac{F(x+h)-F(x)}{h} = \frac{G(x)-F(x)}{h}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(1)')

doc.add_paragraph('Отже:')

# Формула (2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h \approx \frac{G(x)-F(x)}{F'(x)}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(2)')

doc.add_paragraph('Усереднення декількох оцінок:')

# Формула (3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \sum_x \frac{G(x)-F(x)}{F'(x)} / \sum_x 1")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(3)')

doc.add_paragraph('Зважене усереднення з ваговою функцією w(x) = 1/|G\'(x)-F\'(x)|:')

# Формула (6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \sum_x \frac{w(x)[G(x)-F(x)]}{F'(x)} / \sum_x w(x)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(6)')

doc.add_paragraph('Ітерація Ньютона-Рафсона:')

# Формула (7)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h_0 = 0,\quad h_{k+1} = h_k + \sum_x \frac{w(x)[G(x)-F(x+h_k)]}{F'(x+h_k)} / \sum_x w(x)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(7)')

# ============================================================================
# 4.2. Альтернативний висновок (мінімізація L₂ норми)
# ============================================================================

doc.add_heading('4.2. Альтернативний висновок (мінімізація L₂ норми)', level=3)

doc.add_paragraph('Лінійна апроксимація:')

# Формула (8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"F(x+h) \approx F(x) + hF'(x)")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(8)')

doc.add_paragraph('Функція похибки:')

# Формула (9) - перша частина
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E = \sum_x [F(x) + hF'(x) - G(x)]^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(9)')

doc.add_paragraph('Прирівнювання похідної до нуля дає:')

# Формула (10)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \frac{\sum_x F'(x)[G(x)-F(x)]}{\sum_x F'(x)^2}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(10)')

# ============================================================================
# 4.4. Узагальнення на багатовимірний випадок
# ============================================================================

doc.add_heading('4.4. Узагальнення на багатовимірний випадок', level=3)

doc.add_paragraph('Для багатовимірного випадку мінімізуємо:')

# Формула (11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"E = \sum_{x \in R} [F(x+h) - G(x)]^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(11)')

doc.add_paragraph('Лінійна апроксимація в n-вимірах:')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"F(x+h) \approx F(x) + h \frac{\partial F}{\partial x}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(12)')

doc.add_paragraph('Розв\'язок для h:')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"h = \left[ \sum_x \left( \frac{\partial F}{\partial x} \right)^T (G(x)-F(x)) \right] \left[ \sum_x \left( \frac{\partial F}{\partial x} \right)^T \left( \frac{\partial F}{\partial x} \right) \right]^{-1}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(13)')

# ============================================================================
# 5. ЗАСТОСУВАННЯ ДО СТЕРЕОБАЧЕННЯ
# ============================================================================

doc.add_heading('5. Застосування до стереобачення', level=2)

doc.add_paragraph("""
Узагальнений метод реєстрації може бути застосований для отримання інформації про глибину зі стереопар зображень.
Для об'єкта на відстані z від камери 1, положення на площині плівки камери 2 може бути обчислено.
Використовуючи лінійну апроксимацію для корекції глибини Δz:
""")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\Delta z = \sum_x \frac{\partial F}{\partial z} [G-F] / \sum_x \left( \frac{\partial F}{\partial z} \right)^2")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(14)')

doc.add_paragraph("""
де ∂F/∂z обчислюється через правило ланцюга з просторового градієнта інтенсивності та геометрії стереоустановки.
Аналогічно, для невідомих параметрів камери c, оновлення Δc має вигляд:
""")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
math2docx.add_math(p, r"\Delta c = \left[ \sum_x \left( \frac{\partial q}{\partial c} \frac{\partial F}{\partial q} \right)^T (G-F) \right] \left[ \sum_x \left( \frac{\partial q}{\partial c} \frac{\partial F}{\partial q} \right)^T \left( \frac{\partial q}{\partial c} \frac{\partial F}{\partial q} \right) \right]^{-1}")
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_num.add_run('(15)')

# ============================================================================
# ВИСНОВКИ
# ============================================================================

doc.add_heading('Висновки', level=2)

conclusions = """1. Запропонований ітеративний метод реєстрації зображень використовує інформацію про просторовий градієнт інтенсивності для спрямування пошуку, що потребує значно менше порівнянь, ніж методи повного перебору.

2. Метод забезпечує ітерації типу Ньютона-Рафсона, які збігаються квадратично, коли початкова оцінка є достатньо близькою.

3. Метод природним чином узагальнюється на багатовимірний випадок та на довільні лінійні перетворення (обертання, масштабування, зсув).

4. У застосуваннях до стереобачення алгоритм може одночасно визначати глибину об'єктів та параметри камери.

5. Стратегія «від грубого до точного» з використанням смугових фільтрів розширює діапазон збіжності, зберігаючи точність."""
doc.add_paragraph(conclusions)

# ============================================================================
# ЗБЕРЕЖЕННЯ
# ============================================================================

output_file = 'Lucas_Kanade_1981_UKR.docx'
doc.save(output_file)

print("=" * 70)
print(f"✅ Документ '{output_file}' успішно створено українською мовою!")
print("=" * 70)
print("\n📄 Документ містить:")
print("   • Повну назву статті та інформацію про авторів")
print("   • Анотацію (Abstract) українською")
print("   • Вступ (Introduction)")
print("   • Постановку задачі реєстрації зображень")
print("   • Одновимірний випадок методу Лукаса-Канаде (формули 1-7)")
print("   • Альтернативний висновок через мінімізацію L₂ норми (формули 8-10)")
print("   • Узагальнення на багатовимірний випадок (формули 11-13)")
print("   • Застосування до стереозору (формули 14-15)")
print("   • Висновки")